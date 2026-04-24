from fastapi import FastAPI
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import re
import os
import requests
from docx import Document
import tempfile

app = FastAPI(title="AI教案生成系统")

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===================== AI 配置 =====================
DOUBAO_API_KEY = "ark-c923f9c2-0215-4823-b2e0-792b1ed1cbe5-7e5fe"
DOUBAO_URL = "https://ark.cn-beijing.volces.com/api/v3/responses"
DOUBAO_MODEL = "doubao-seed-2-0-pro-260215"

# ===================== 读取指令 =====================
def get_prompt(topic):
    try:
        with open("AI 生成教案专用指令.txt", "r", encoding="utf-8") as f:
            prompt = f.read()
        return prompt.strip() + f"\n课程主题：{topic}"
    except:
        return f"请生成职业院校专业课教案，格式&数字&内容，课程主题：{topic}"

# ===================== 调用AI =====================
def ai_generate(topic):
    prompt = get_prompt(topic)
    headers = {
        "Authorization": f"Bearer {DOUBAO_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": DOUBAO_MODEL,
        "input": [{"role": "user", "content": [{"type": "input_text", "text": prompt}]}]
    }
    try:
        resp = requests.post(DOUBAO_URL, headers=headers, json=data, timeout=180)
        res = resp.json()
        for item in res.get("output", []):
            if item.get("role") == "assistant":
                for c in item.get("content", []):
                    if c.get("type") == "output_text":
                        return c.get("text", "")
    except Exception as e:
        print("AI错误：", e)
    return ""

# ===================== 解析格式 =====================
def parse_content(text):
    data = {}
    matches = re.findall(r"&(\d+)&(.*?)&\1&", text, re.DOTALL)
    for num, val in matches:
        data[int(num)] = val.strip()
    return data

# ===================== 生成Word =====================
def make_docx(data):
    doc = Document("模板.docx")
    # 替换段落
    for p in doc.paragraphs:
        for k, v in data.items():
            p.text = p.text.replace(f"&{k}&", v)
    # 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in data.items():
                        p.text = p.text.replace(f"&{k}&", v)
    tmp = tempfile.mktemp(".docx")
    doc.save(tmp)
    return tmp

# ===================== 接口 =====================
@app.get("/")
def index():
    with open("web.html", "r", encoding="utf-8") as f:
        return HTMLResponse(f.read())

@app.get("/generate")
def generate(topic: str):
    text = ai_generate(topic)
    if not text:
        return {"code": 1, "msg": "AI生成失败"}
    data = parse_content(text)
    if not data:
        return {"code": 1, "msg": "解析失败，请检查AI返回格式"}
    docx_file = make_docx(data)
    return FileResponse(
        path=docx_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{topic}教案.docx"
    )

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)